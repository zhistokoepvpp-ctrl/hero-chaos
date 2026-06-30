import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    # styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.rstrip()

        # skip separators
        if line.strip() == '---':
            continue

        # headers
        if line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        # code blocks
        elif line.startswith('```'):
            continue
        elif line.startswith('| '):
            # simple table handling - skip for now
            p = doc.add_paragraph(line)
            p.style = doc.styles['Normal']
            p.paragraph_format.space_after = Pt(2)
        elif line.strip().startswith('|') and line.strip().endswith('|'):
            continue
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            p = doc.add_paragraph(line, style='List Number')
            p.paragraph_format.space_after = Pt(2)
        elif line.strip() == '':
            doc.add_paragraph('')
        else:
            # inline code in text
            text = line
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(2)

    doc.save(docx_path)
    print(f'Saved: {docx_path}')

if __name__ == '__main__':
    md_to_docx(sys.argv[1], sys.argv[2])
